"""Generate PRODUCT.docx from structured content."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, header_color="1a365d"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f0f4f8")

    return table


def build_document():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Styles --
    style = doc.styles["Title"]
    style.font.size = Pt(28)
    style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    for level in range(1, 4):
        s = doc.styles[f"Heading {level}"]
        s.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # ========== COVER ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ForensicInvestigator")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Document Produit")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(
        "Plateforme d'aide à l'investigation criminelle\n"
        "augmentée par l'intelligence artificielle"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    for _ in range(8):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Domaine : Investigation criminelle / Analyse forensique\n"
        "Utilisateurs cibles : Enquêteurs PJ, analystes criminels, magistrats\n"
        "Interface : Français\n\n"
        "Mars 2026"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ========== TABLE DES MATIERES ==========
    doc.add_heading("Table des matières", level=1)
    toc_items = [
        "1. Vue d'ensemble",
        "2. Proposition de valeur",
        "3. Architecture technique",
        "4. Modules fonctionnels",
        "   4.1  Gestion des affaires",
        "   4.2  Graphe de connaissances interactif",
        "   4.3  Analyse de graphe avancée",
        "   4.4  Raisonnement hiérarchique (HRM)",
        "   4.5  Assistant conversationnel",
        "   4.6  N4L (Narrative 4 Logic)",
        "   4.7  Détection d'anomalies",
        "   4.8  Simulation What-If",
        "   4.9  Analyse cross-affaires",
        "   4.10 Recherche hybride",
        "   4.11 Cartographie géographique",
        "   4.12 Réseau social",
        "   4.13 Investigation guidée (PEACE)",
        "   4.14 Carnet d'analyse",
        "   4.15 Configuration des prompts",
        "5. Données de démonstration",
        "6. Déploiement",
        "7. Métriques du code",
        "8. Sécurité",
        "9. Roadmap",
        "10. Points forts différenciants",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ========== 1. VUE D'ENSEMBLE ==========
    doc.add_heading("1. Vue d'ensemble", level=1)
    doc.add_paragraph(
        "ForensicInvestigator est une plateforme d'aide à l'investigation criminelle "
        "augmentée par l'intelligence artificielle. Elle centralise la gestion des affaires, "
        "l'analyse des preuves, la visualisation des réseaux relationnels et le raisonnement "
        "déductif automatisé pour accompagner les enquêteurs dans la résolution d'affaires complexes."
    )
    doc.add_paragraph(
        "La plateforme intègre un moteur de graphe de connaissances (SSTorytime), un langage "
        "formel de description narrative (N4L), et un système de raisonnement hiérarchique (HRM) "
        "alimenté par un LLM pour offrir une approche multi-dimensionnelle de l'analyse forensique."
    )

    # ========== 2. PROPOSITION DE VALEUR ==========
    doc.add_heading("2. Proposition de valeur", level=1)
    doc.add_paragraph(
        "ForensicInvestigator répond à six problèmes majeurs rencontrés dans les enquêtes complexes :"
    )
    add_styled_table(doc,
        ["Problème", "Solution ForensicInvestigator"],
        [
            [
                "Les affaires complexes impliquent des centaines d'entités et relations difficiles à appréhender",
                "Graphe de connaissances interactif avec algorithmes d'analyse avancés (clusters, centralité, chemins)",
            ],
            [
                "Les contradictions dans les témoignages passent inaperçues",
                "Détection automatique de contradictions par IA avec localisation précise",
            ],
            [
                "Les liens entre affaires distinctes sont rarement identifiés",
                "Analyse cross-affaires avec matching sémantique et scoring de similarité",
            ],
            [
                "Le raisonnement déductif est long et sujet aux biais cognitifs",
                "Raisonnement hiérarchique multi-étapes (HRM) avec scoring de confiance",
            ],
            [
                "Les anomalies temporelles et comportementales sont noyées dans la masse",
                "Détection statistique d'anomalies (Z-score, MAD) avec alertes par sévérité",
            ],
            [
                "Les hypothèses sont rarement testées formellement",
                "Simulation What-If avec propagation des implications sans altérer les données",
            ],
        ],
    )

    # ========== 3. ARCHITECTURE TECHNIQUE ==========
    doc.add_heading("3. Architecture technique", level=1)

    doc.add_heading("Stack technologique", level=2)
    add_styled_table(doc,
        ["Composant", "Technologie", "Port"],
        [
            ["Backend API", "Go 1.24 (stdlib net/http)", "8082"],
            ["Moteur de raisonnement HRM", "Python / FastAPI / llama.cpp", "8081"],
            ["Service d'embeddings", "multilingual-e5-base sur SPARK (768 dim)", "8002"],
            ["Frontend", "JavaScript vanilla (modulaire)", "—"],
            ["Visualisation graphe", "vis-network", "—"],
            ["Cartographie", "Leaflet", "—"],
            ["LLM distant", "vLLM sur SPARK GB10 (Qwen3.8-27B-FP8, 128k ctx)", "8001"],
            ["Analyse de graphe", "SSTorytime (package Go)", "—"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Architecture des services", level=2)
    arch_text = (
        "L'application repose sur trois services indépendants orchestrés par des scripts shell :\n\n"
        "1. Backend Go (port 8082) — Serveur HTTP principal avec 80+ endpoints API, "
        "11 services métier (CaseService, N4LService, GraphAnalyzerService, SearchService, "
        "AnomalyService, ScenarioService, NotebookService, OllamaService, HRMService, "
        "N4LGeneratorService, ConfigService), et serveur de fichiers statiques.\n\n"
        "2. Serveur HRM (port 8081) — Service Python FastAPI dédié au raisonnement déductif. "
        "Deux moteurs : HRM Engine (règles) et HRM Sapient (vLLM). "
        "Raisonnement en 3 phases (planification, exécution, synthèse) avec streaming SSE.\n\n"
        "3. Service d'embeddings (port 8002, SPARK) — API compatible OpenAI servant multilingual-e5-base "
        "(768 dimensions). Fournit les vecteurs du moteur de recherche hybride, avec cache local."
    )
    doc.add_paragraph(arch_text)

    # ========== 4. MODULES FONCTIONNELS ==========
    doc.add_heading("4. Modules fonctionnels", level=1)
    doc.add_paragraph(
        "L'application frontend est une SPA (Single Page Application) composée de 22 modules "
        "JavaScript indépendants, chacun associé à un ou plusieurs modules CSS dédiés."
    )

    # 4.1
    doc.add_heading("4.1 Gestion des affaires", level=2)
    doc.add_paragraph(
        "Création, consultation et mise à jour des dossiers d'enquête. Chaque affaire centralise :"
    )
    for item in [
        "Entités : personnes, lieux, objets, organisations, documents, événements",
        "Preuves : physiques, testimoniales, documentaires, numériques, forensiques",
        "Relations typées entre entités (connaît, possède, était_présent, etc.)",
        "Chronologie des événements avec horodatage précis",
        "Hypothèses d'investigation avec scoring de confiance",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 4.2
    doc.add_heading("4.2 Graphe de connaissances interactif", level=2)
    doc.add_paragraph(
        "Visualisation réseau des entités et relations d'une affaire via la bibliothèque vis-network. "
        "Navigation interactive avec zoom, drag, filtrage par type d'entité, menu contextuel "
        "sur les nœuds, coloration par type et mise en évidence des chemins et clusters."
    )

    # 4.3
    doc.add_heading("4.3 Analyse de graphe avancée (SSTorytime)", level=2)
    doc.add_paragraph(
        "Algorithmes de théorie des graphes appliqués à l'investigation, intégrés via le package SSTorytime :"
    )
    add_styled_table(doc,
        ["Algorithme", "Usage investigatif"],
        [
            ["Détection de clusters", "Identifier les groupes organisés"],
            ["Plus court chemin", "Tracer les chaînes relationnelles"],
            ["Centralité eigenvector", "Repérer les acteurs influents"],
            ["Centralité betweenness", "Identifier les intermédiaires clés"],
            ["Super-nœuds", "Détecter les entités à haute connectivité"],
            ["Cône d'expansion", "Explorer le voisinage d'une entité"],
            ["Analyse de densité", "Évaluer la couverture de l'enquête"],
            ["Patterns temporels", "Détecter les rythmes d'activité"],
            ["Propagation contra", "Tracer la diffusion des contradictions"],
            ["Organisation en couches", "Hiérarchiser le réseau"],
        ],
    )

    # 4.4
    doc.add_heading("4.4 Raisonnement hiérarchique (HRM)", level=2)
    doc.add_paragraph(
        "Moteur de raisonnement déductif multi-étapes alimenté par un LLM (Qwen3.8-27B-FP8 via vLLM, 128K contexte). "
        "Le processus se déroule en trois phases :"
    )
    for phase, desc in [
        ("Planification", "Décomposition de la question en sous-tâches : analyser les preuves, identifier les acteurs, construire la chronologie, évaluer les hypothèses."),
        ("Exécution", "Résolution séquentielle de chaque sous-tâche avec le contexte complet de l'affaire."),
        ("Synthèse", "Consolidation en conclusion structurée avec niveau de confiance, conclusions alternatives, avertissements et recommandations."),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{phase} — ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "L'interface affiche la progression en temps réel via streaming SSE, "
        "avec indicateurs visuels par phase."
    )

    # 4.5
    doc.add_heading("4.5 Assistant conversationnel (Chat IA)", level=2)
    doc.add_paragraph(
        "Interface de dialogue naturel pour interroger les données d'une affaire. "
        "Streaming des réponses en temps réel, contexte automatique (entités, preuves, relations), "
        "historique de conversation et rendu Markdown."
    )

    # 4.6
    doc.add_heading("4.6 N4L (Narrative 4 Logic)", level=2)
    doc.add_paragraph(
        "Langage formel de description narrative pour encoder les relations d'une affaire. "
        "Exemple : « Pierre Dupont -[connaît]-> Marie Martin ». "
        "L'application intègre un éditeur avec coloration syntaxique, un parseur bidirectionnel "
        "(N4L ↔ Graphe), un générateur automatique depuis le texte libre, et le support des "
        "contextes, alias, groupes et relations chaînées."
    )

    # 4.7
    doc.add_heading("4.7 Détection d'anomalies", level=2)
    doc.add_paragraph(
        "Analyse statistique automatisée pour identifier les points atypiques dans les données "
        "d'une affaire. Utilise les méthodes Z-score et Z-score modifié (MAD) pour la robustesse. "
        "Détecte les anomalies temporelles (horaires inhabituels), comportementales (changements de pattern) "
        "et relationnelles (connexions inattendues). Système d'alertes avec niveaux de sévérité "
        "et tableau de bord statistique."
    )

    # 4.8
    doc.add_heading("4.8 Simulation What-If (Scénarios)", level=2)
    doc.add_paragraph(
        "Permet de tester des hypothèses sans altérer les données réelles de l'affaire. "
        "Création de scénarios hypothétiques, propagation automatique des implications, "
        "scoring de plausibilité et comparaison de scénarios alternatifs."
    )

    # 4.9
    doc.add_heading("4.9 Analyse cross-affaires", level=2)
    doc.add_paragraph(
        "Détection automatique de liens entre affaires distinctes. Matching d'entités par nom "
        "et attributs, détection de patterns récurrents (modus operandi), analyse des correspondances "
        "avec déduplication stricte et scoring de similarité."
    )

    # 4.10
    doc.add_heading("4.10 Recherche hybride", level=2)
    doc.add_paragraph(
        "Moteur de recherche combinant BM25 (recherche plein texte classique par mots-clés) "
        "et recherche sémantique (embeddings multilingual-e5-base pour capturer le sens). "
        "Classement fusionné pour des résultats pertinents même avec des formulations différentes."
    )

    # 4.11
    doc.add_heading("4.11 Cartographie géographique", level=2)
    doc.add_paragraph(
        "Visualisation spatiale via Leaflet. Positionnement des lieux sur carte, "
        "analyse des patterns géographiques et corrélation spatiale des événements."
    )

    # 4.12
    doc.add_heading("4.12 Réseau social", level=2)
    doc.add_paragraph(
        "Cartographie relationnelle centrée sur les personnes. Visualisation des liens "
        "interpersonnels, analyse de la structure sociale et identification des réseaux d'influence."
    )

    # 4.13
    doc.add_heading("4.13 Investigation guidée (méthode PEACE)", level=2)
    doc.add_paragraph(
        "Workflow structuré pour les interrogatoires et la collecte d'informations. "
        "Génération automatique de questions pertinentes selon la méthode PEACE, "
        "cadre méthodologique et traçabilité des réponses."
    )

    # 4.14
    doc.add_heading("4.14 Carnet d'analyse (Notebook)", level=2)
    doc.add_paragraph(
        "Centralisation des analyses IA. Sauvegarde des résultats d'analyse, "
        "système de tags et favoris, épinglage des notes importantes, persistance en JSON."
    )

    # 4.15
    doc.add_heading("4.15 Configuration des prompts", level=2)
    doc.add_paragraph(
        "Interface d'administration pour personnaliser les 9 catégories de prompts IA : "
        "analyse de cas, génération d'hypothèses, détection de contradictions, génération de questions, "
        "analyse d'hypothèse, analyse de chemin, chat, analyse cross-affaires, conversion N4L. "
        "Rechargement à chaud sans redémarrage du service."
    )

    # ========== 5. DONNEES DE DEMO ==========
    doc.add_heading("5. Données de démonstration", level=1)
    doc.add_paragraph(
        "Six affaires pré-chargées couvrant différents types d'infractions :"
    )
    add_styled_table(doc,
        ["Affaire", "Type", "Complexité"],
        [
            ["Affaire Victor Moreau", "Homicide par empoisonnement", "Élevée"],
            ["Affaire Disparition", "Disparition de personne", "Moyenne"],
            ["Affaire Fraude", "Fraude financière", "Moyenne"],
            ["Affaire Cambriolage", "Vol avec effraction", "Standard"],
            ["Affaire Incendie", "Incendie criminel", "Moyenne"],
            ["Affaire Trafic Art", "Trafic d'œuvres d'art", "Élevée"],
        ],
    )

    # ========== 6. DEPLOIEMENT ==========
    doc.add_heading("6. Déploiement", level=1)

    doc.add_heading("Développement local", level=2)
    doc.add_paragraph(
        "Deux scripts shell orchestrent le démarrage et l'arrêt des trois services : "
        "start_services.sh et stop_services.sh. Les logs sont centralisés dans le dossier logs/."
    )

    doc.add_heading("Production (Ubuntu + systemd)", level=2)
    doc.add_paragraph(
        "Un script d'installation (scripts/install.sh) configure deux unités systemd : "
        "forensicinvestigator (Go, port 8082) et forensicinvestigator-hrm (Python, port 8081). "
        "Le LLM et les embeddings sont fournis par le serveur SPARK (ports 8001 et 8002). "
        "Une configuration Nginx (scripts/nginx_forensic.conf) assure le reverse proxy "
        "avec support HTTPS via Let's Encrypt."
    )

    doc.add_heading("Prérequis", level=2)
    for item in [
        "Go 1.24+",
        "Python 3.12+ avec venv",
        "Accès au serveur LLM distant llama.cpp (ou Ollama local en fallback)",
        "~4 Go RAM minimum",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ========== 7. METRIQUES ==========
    doc.add_heading("7. Métriques du code", level=1)
    add_styled_table(doc,
        ["Composant", "Volume"],
        [
            ["Frontend JavaScript", "~27 000 lignes"],
            ["Styles CSS", "~21 000 lignes"],
            ["Backend Go", "~15 000 lignes"],
            ["Services Python", "~2 500 lignes"],
            ["Modules JS", "22 modules"],
            ["Modules CSS", "30 fichiers"],
            ["Endpoints API", "80+ routes"],
            ["Services Go", "11 services"],
            ["Total", "~88 000 lignes"],
        ],
    )

    # ========== 8. SECURITE ==========
    doc.add_heading("8. Sécurité", level=1)
    for item in [
        "Authentification par mot de passe à l'entrée de l'application",
        "Validation des entrées utilisateur côté serveur",
        "Thread-safety via sync.RWMutex sur les structures de données partagées",
        "Variables d'environnement pour les secrets (pas de hardcoding)",
        "Configuration Nginx pour HTTPS en production",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ========== 9. ROADMAP ==========
    doc.add_heading("9. Roadmap potentielle", level=1)
    add_styled_table(doc,
        ["Priorité", "Fonctionnalité"],
        [
            ["Haute", "Persistance en base de données (PostgreSQL)"],
            ["Haute", "Authentification multi-utilisateurs avec rôles"],
            ["Moyenne", "Export PDF des rapports d'analyse"],
            ["Moyenne", "Intégration avec les fichiers de police existants"],
            ["Moyenne", "Mode hors-ligne avec LLM local"],
            ["Basse", "Application mobile pour le terrain"],
            ["Basse", "API publique pour intégrations tierces"],
        ],
    )

    # ========== 10. POINTS FORTS ==========
    doc.add_heading("10. Points forts différenciants", level=1)

    differenciants = [
        ("Raisonnement déductif structuré",
         "Le HRM décompose les questions complexes en sous-tâches, contrairement aux chatbots qui répondent en un seul passage."),
        ("Langage formel N4L",
         "Représentation rigoureuse des relations, exportable et réutilisable entre affaires."),
        ("Analyse de graphe native",
         "Algorithmes de théorie des graphes (10 algorithmes) directement intégrés via SSTorytime."),
        ("Détection d'anomalies statistique",
         "Approche quantitative (Z-score, MAD) complémentaire au raisonnement qualitatif de l'IA."),
        ("Simulation sans altération",
         "Les scénarios What-If n'impactent jamais les données réelles de l'affaire."),
        ("Interface temps réel",
         "Streaming SSE pour un feedback immédiat pendant les analyses longues (HRM, Chat)."),
        ("Architecture modulaire",
         "22 modules frontend indépendants, 11 services backend découplés, 3 services déployables séparément."),
    ]

    for i, (title, desc) in enumerate(differenciants, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title} — ")
        run.bold = True
        p.add_run(desc)

    # ========== SAVE ==========
    output = "/Users/ilan/_INFOSTRATES/_AI/SSTorytime-1/src/ForensicInvestigator/PRODUCT.docx"
    doc.save(output)
    print(f"Document saved: {output}")


if __name__ == "__main__":
    build_document()
